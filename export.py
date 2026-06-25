"""
Fonctions d'export du transcript et du résumé en TXT et DOCX.
"""
import io
from datetime import datetime
from typing import Union
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Helpers ───────────────────────────────────────────────────────────────────

def _now_str() -> str:
    return datetime.now().strftime("%d/%m/%Y à %H:%M")


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _set_doc_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)


# ── Transcript ─────────────────────────────────────────────────────────────────

def transcript_to_txt(transcript: str) -> bytes:
    header = f"TRANSCRIPT DE RÉUNION\nGénéré le {_now_str()}\n{'─' * 60}\n\n"
    return (header + transcript).encode("utf-8")


def transcript_to_docx(transcript: str) -> bytes:
    doc = Document()
    _set_doc_margins(doc)

    _add_heading(doc, "Transcript de réunion", level=1)
    doc.add_paragraph(f"Généré le {_now_str()}")
    doc.add_paragraph("")

    for paragraph in transcript.split("\n"):
        p = paragraph.strip()
        if p:
            doc.add_paragraph(p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Résumé ────────────────────────────────────────────────────────────────────

def summary_to_txt(summary: Union[dict, str]) -> bytes:
    lines = []

    if isinstance(summary, str):
        lines.append("RÉSUMÉ DE RÉUNION")
        lines.append(f"Généré le {_now_str()}")
        lines.append("─" * 60)
        lines.append("")
        lines.append(summary)
        return "\n".join(lines).encode("utf-8")

    titre = summary.get("titre_reunion") or "Réunion"
    lines.append(f"RÉSUMÉ — {titre.upper()}")
    lines.append(f"Généré le {_now_str()}")

    if summary.get("date"):
        lines.append(f"Date : {summary['date']}")
    if summary.get("duree_estimee"):
        lines.append(f"Durée estimée : {summary['duree_estimee']}")
    lines.append("─" * 60)
    lines.append("")

    participants = summary.get("participants") or []
    if participants:
        lines.append("PARTICIPANTS")
        lines.append(", ".join(participants))
        lines.append("")

    ordre = summary.get("ordre_du_jour") or []
    if ordre:
        lines.append("ORDRE DU JOUR")
        for item in ordre:
            lines.append(f"  • {item}")
        lines.append("")

    points = summary.get("points_cles") or []
    if points:
        lines.append("POINTS CLÉS")
        for point in points:
            lines.append(f"  • {point}")
        lines.append("")

    decisions = summary.get("decisions") or []
    if decisions:
        lines.append("DÉCISIONS")
        for d in decisions:
            lines.append(f"  • {d.get('decision', '')}")
            if d.get("contexte"):
                lines.append(f"    → {d['contexte']}")
        lines.append("")

    actions = summary.get("actions") or []
    if actions:
        lines.append("ACTIONS")
        for a in actions:
            action_line = f"  • [{a.get('priorite', '').upper()}] {a.get('action', '')}"
            meta = []
            if a.get("responsable"):
                meta.append(a["responsable"])
            if a.get("deadline"):
                meta.append(f"avant le {a['deadline']}")
            if meta:
                action_line += f"  ({', '.join(meta)})"
            lines.append(action_line)
        lines.append("")

    prochaine = summary.get("prochaine_reunion")
    if prochaine:
        lines.append("PROCHAINE RÉUNION")
        lines.append(f"  {prochaine}")

    return "\n".join(lines).encode("utf-8")


def summary_to_docx(summary: Union[dict, str]) -> bytes:
    doc = Document()
    _set_doc_margins(doc)

    if isinstance(summary, str):
        _add_heading(doc, "Résumé de réunion", level=1)
        doc.add_paragraph(f"Généré le {_now_str()}")
        doc.add_paragraph("")
        doc.add_paragraph(summary)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    titre = summary.get("titre_reunion") or "Résumé de réunion"
    _add_heading(doc, titre, level=1)

    meta_parts = [f"Généré le {_now_str()}"]
    if summary.get("date"):
        meta_parts.append(f"Date : {summary['date']}")
    if summary.get("duree_estimee"):
        meta_parts.append(f"Durée estimée : {summary['duree_estimee']}")
    doc.add_paragraph("  ·  ".join(meta_parts))

    participants = summary.get("participants") or []
    if participants:
        _add_heading(doc, "Participants", level=2)
        doc.add_paragraph(", ".join(participants))

    ordre = summary.get("ordre_du_jour") or []
    if ordre:
        _add_heading(doc, "Ordre du jour", level=2)
        for item in ordre:
            _add_bullet(doc, item)

    points = summary.get("points_cles") or []
    if points:
        _add_heading(doc, "Points clés", level=2)
        for point in points:
            _add_bullet(doc, point)

    decisions = summary.get("decisions") or []
    if decisions:
        _add_heading(doc, "Décisions", level=2)
        for d in decisions:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(d.get("decision", ""))
            run.bold = True
            if d.get("contexte"):
                p.add_run(f" — {d['contexte']}")

    actions = summary.get("actions") or []
    if actions:
        _add_heading(doc, "Actions", level=2)
        priority_label = {"haute": "[HAUTE]", "moyenne": "[MOY.]", "basse": "[BASSE]"}
        for a in actions:
            priorite = (a.get("priorite") or "basse").lower()
            p = doc.add_paragraph(style="List Bullet")
            label = priority_label.get(priorite, "")
            p.add_run(f"{label} ").bold = True
            p.add_run(a.get("action", ""))
            meta = []
            if a.get("responsable"):
                meta.append(a["responsable"])
            if a.get("deadline"):
                meta.append(f"avant le {a['deadline']}")
            if meta:
                run = p.add_run(f"  ({', '.join(meta)})")
                run.italic = True

    prochaine = summary.get("prochaine_reunion")
    if prochaine:
        _add_heading(doc, "Prochaine réunion", level=2)
        doc.add_paragraph(prochaine)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
